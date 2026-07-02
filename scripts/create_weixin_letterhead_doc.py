from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGO = Path("/Users/rushi/Downloads/New Logo - Blip and Text.png")
STAMP = ROOT / "generated-documents" / "whats-on-company-stamp-blue-transparent.png"
OUT = ROOT / "generated-documents" / "weixin_official_account_application_whats_on_campus_stamped.docx"

GOLD = "FADC4A"
BLACK = "000000"
WHITE = "FFFFFF"
MUTED = "666666"
RULE = "D9DDE3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "FFFFFF", size: str = "0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil" if size == "0" else "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_table_grid(table, widths: list[int]) -> None:
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_border_bottom(paragraph, color: str, size: str = "12", space: str = "3") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Number", "List Paragraph"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def build_letterhead_header(header) -> None:
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].text = ""
        header.paragraphs[0].paragraph_format.space_after = Pt(0)

    table = header.add_table(rows=2, cols=1, width=Inches(6.5))
    set_table_width(table, 9360, 0)
    set_table_grid(table, [9360])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_cell = table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(logo_cell, BLACK)
    set_cell_borders(logo_cell)
    set_cell_margins(logo_cell, top=90, bottom=90, start=120, end=120)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(0)
    logo_run = logo_para.add_run()
    logo_run.add_picture(str(LOGO), width=Inches(2.55))

    gold_cell = table.cell(1, 0)
    set_cell_shading(gold_cell, GOLD)
    set_cell_borders(gold_cell, GOLD, "4")
    set_cell_margins(gold_cell, top=18, bottom=18, start=0, end=0)
    gold_cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def build_running_header(header) -> None:
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].text = ""
        header.paragraphs[0].paragraph_format.space_after = Pt(0)

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("What's On! Campus Pty Ltd | Weixin Official Account Application")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_border_bottom(p, GOLD, size="8", space="2")


def build_header(section) -> None:
    build_letterhead_header(section.first_page_header)
    build_running_header(section.header)


def build_footer_part(footer) -> None:
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_border_bottom(p, RULE, size="4", space="2")
    r = p.add_run("What's On! Campus Pty Ltd | Weixin Official Account Application")
    set_run_font(r, size=8.5, color=MUTED)


def build_footer(section) -> None:
    build_footer_part(section.first_page_footer)
    build_footer_part(section.footer)


def add_title_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("WEIXIN OFFICIAL ACCOUNT REGISTRATION AND VERIFICATION APPLICATION FORM")
    set_run_font(r, size=14, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("(International Applicant)")
    set_run_font(r, size=11, color=MUTED, bold=True)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(14)
    set_paragraph_border_bottom(rule, GOLD, size="18", space="2")


def add_opening(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    parts = [
        ("The Applicant(entity name) ", False),
        ("What's On! Campus Pty Ltd", True),
        (" hereby irrevocably applies for registration and verification Weixin "
         "(Official Accounts/Mini Programs/Open Platform) (Original ID) ", False),
        ("gh_119addaec551", True),
        ("(the “Account”), and agrees with the followings:", False),
    ]
    for text, bold in parts:
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK, bold=bold)


def add_numbered_clause(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)


def add_body(doc: Document) -> None:
    clauses = [
        "The Applicant represents and warrants that it has full and complete capability and authority to authorize the Contact Person who is designated by the Applicant and submitted to Tencent in the process of registration and verification of the Account (including any changes to the Contact Person which may be submitted by the Applicant from time to time via designated Account channels) to irrevocably apply for registration, verification, operation, maintenance and management of the Account on behalf of the Applicant, to pay application fee by WeChat Pay, credit card or debit card of the Contact Person for the Applicant and that the Contact Person has full and complete capability and authority to perform the foregoing authorization. Any and all obligations and liabilities arising from or related to the Account shall be assumed by the Applicant and Contact Person.",
        "The Applicant hereby agrees that, when applying for registration and verification of the Account, the Applicant shall fill in its name and confirm by signing this application form. After successful Account verification, verified Applicants shall have the right to use the Account. Any and all rights and obligations arising from the Account as from the date of registration shall be assumed by the Applicant. The verified Applicants shall be entitled to any and all earnings and permissions obtained by the Account, and all operation activities must be carried out by the Applicant.",
        "The Applicant hereby acknowledges and represents that all registration and verification information that are submitted to Tencent are true, accurate and valid, and that it irrevocably authorizes Tencent and any third-party verification organization delegated by Tencent to review and verify the submitted materials. The Applicant acknowledges and agrees that the verification service fee paid will not be refunded regardless of the verification result, withdrawal of application by the Applicant or any other factors.",
        "The Applicant acknowledges and agrees that content maintenance, development maintenance and operation management of the Account shall comply with the applicable laws and regulations of the People’s Republic of China and local rules, as well as relevant agreements and rules of Tencent (collectively, the “Tencent Agreements”) , including but not limited to Tencent Service Agreement, Weixin Official Accounts Platform Service Agreement, Weixin Official Accounts Platform Authentication Service Agreement, Weixin Mini-Program Platform – Terms of Service, Weixin Open Platform Developer Service Agreement. The Applicant shall solely assume all the liabilities if it violates any of the foregoing commitments.",
        "The Applicant acknowledges and agrees that this registration and verification process is limited to reviewing the authenticity and legality of the materials submitted by the Applicant. Whether certain features and permissions can be granted, or whether the Account can be created shall be subject to specific rules formulated by the platform, which is independent of the verification result. Tencent may, at its sole discretion, accept or reject the application for registration and verification of the Account.",
    ]
    for clause in clauses:
        add_numbered_clause(doc, clause)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(
        "The Applicant hereby confirms that there is no objection to the above registration and verification "
        "information and to this application form. This application form shall remain in full force and effect, "
        "provided that no changes occur to the Applicant."
    )
    set_run_font(r, size=11, color=BLACK)


def add_signature_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Sign by authorized representative of the Applicant:")
    set_run_font(r, size=11, color=BLACK, bold=True)

    table = doc.add_table(rows=2, cols=2)
    set_table_width(table, 9360, 120)
    set_table_grid(table, [5148, 4212])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, RULE, "4")
            set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.cell(0, 0)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Authorized representative signature")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_border_bottom(p, "9CA3AF", size="6", space="1")

    right = table.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if STAMP.exists():
        stamp_run = p.add_run()
        stamp_run.add_picture(str(STAMP), width=Inches(2.7))
    else:
        set_cell_shading(right, "FFF8D6")
        p.paragraph_format.space_before = Pt(20)
        r = p.add_run("Please affix with the company seal")
        set_run_font(r, size=10, color=BLACK, bold=True)

    date_cell = table.cell(1, 0)
    p = date_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Date: ")
    set_run_font(r, size=11, color=BLACK, bold=True)
    r = p.add_run("01/07/2026")
    set_run_font(r, size=11, color=BLACK)

    table.cell(1, 1).text = ""
    p = table.cell(1, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    build_header(doc.sections[0])
    build_footer(doc.sections[0])
    add_title_block(doc)
    add_opening(doc)
    add_body(doc)
    add_signature_block(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
