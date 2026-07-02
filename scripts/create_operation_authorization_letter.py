from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from create_weixin_letterhead_doc import (
    BLACK,
    GOLD,
    MUTED,
    RULE,
    STAMP,
    build_footer,
    build_header,
    set_cell_borders,
    set_cell_margins,
    set_cell_shading,
    set_paragraph_border_bottom,
    set_run_font,
    set_table_grid,
    set_table_width,
    style_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "generated-documents" / "operation_authorization_letter_whats_on_campus_stamped.docx"


def add_title_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("OPERATION AUTHORIZATION LETTER")
    set_run_font(r, size=15, color=BLACK, bold=True)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(16)
    set_paragraph_border_bottom(rule, GOLD, size="18", space="2")


def add_body(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Shenzhen Tencent Computer System Co., Ltd.:")
    set_run_font(r, size=11, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.16

    parts = [
        ("The applicant ", False),
        ("What's On! Campus Pty Ltd", True),
        (" hereby applies for the registration and authentication service of WeChat "
         "(public account/mini program/open platform) account (original ID) ", False),
        ("gh_119addaec551", True),
        (" account nickname ", False),
        ("Rush", True),
        (" (hereinafter referred to as \"account\"), and agrees to authorize the appointment of designated administrator ", False),
        ("Rushi Kaushik Vyas", True),
        (" (ID number: ", False),
        ("24639765", True),
        (" Phone number: ", False),
        ("+610416876876", True),
        (") As the administrator of the WeChat account, this administrator can log in to the public account by himself "
         "after the applicant confirms the authorization, and perform operations such as group posting and authorize him "
         "to be responsible for the content maintenance and operation of the account management. The subject of this "
         "application promises that the authentication information submitted to Tencent is true and correct, and it "
         "authorizes Tencent and its authorized third-party review agency to screen and verify the submitted information. "
         "At the same time, the maintenance and operation management of WeChat account content comply with the relevant "
         "provisions of national laws, regulations, policies and WeChat public platform service agreements. If the above "
         "undertaking is violated, the applicant and the administrator shall bear joint and several liabilities. This "
         "authorization letter can only guarantee that this account will authorize the administrator at the time of "
         "authentication performed on the date ", False),
        ("01/07/2026", True),
        (" . Subsequent operations such as changing the administrator of this account or changing the account name cause "
         "the relevant information to be inconsistent with the information written in this authorization, and the back-end "
         "registration of the WeChat public platform shall prevail, and this authorization will automatically become "
         "invalid. The applicant has no objection to the above authorization content confirmation.", False),
    ]
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, size=11, color=BLACK, bold=bold)


def add_signature_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Sign by authorized representative of the Applicant:")
    set_run_font(r, size=11, color=BLACK, bold=True)

    table = doc.add_table(rows=2, cols=2)
    set_table_width(table, 9360, 120)
    set_table_grid(table, [5148, 4212])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, RULE, "4")
            set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.cell(0, 0)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Authorized representative signature")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_border_bottom(p, "9CA3AF", size="6", space="1")
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Rushi Kaushik Vyas")
    set_run_font(r, size=10, color=BLACK)

    right = table.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if STAMP.exists():
        p.add_run().add_picture(str(STAMP), width=Inches(2.7))
    else:
        set_cell_shading(right, "FFF8D6")
        r = p.add_run("Please affix with the company seal")
        set_run_font(r, size=10, color=BLACK, bold=True)

    date_cell = table.cell(1, 0)
    p = date_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("DATE: ")
    set_run_font(r, size=11, color=BLACK, bold=True)
    r = p.add_run("01/07/2026")
    set_run_font(r, size=11, color=BLACK)

    table.cell(1, 1).text = ""
    p = table.cell(1, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    build_header(doc.sections[0])
    build_footer(doc.sections[0])
    add_title_block(doc)
    add_body(doc)
    add_signature_block(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
